"""Reusable, UI-independent Word document generation for Tender Sathi.

The web app streams the returned bytes directly to the browser. No generated
files, images, signatures, stamps, or PDFs are written by this module.
"""

from __future__ import annotations

import io
import re
from decimal import Decimal
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from docx.oxml import OxmlElement
from lxml import etree

from format_utils import format_percentage


class BidDocumentGenerator:
    """Core placeholder replacement and legacy bid template generation."""

    def __init__(self, base_path, create_dirs=True):
        self.base_path = Path(base_path)
        if create_dirs:
            self.ensure_dirs()

    def ensure_dirs(self):
        # Kept for compatibility with the existing test suite and template
        # layout. The web application never writes generated output here.
        for directory in ("templates", "output", "assets"):
            (self.base_path / directory).mkdir(parents=True, exist_ok=True)

    @staticmethod
    def is_empty_value(value):
        return value is None or str(value).strip() == ""

    def replace_all_in_paragraph(self, paragraph, placeholders):
        for run in paragraph.runs:
            for key, value in placeholders.items():
                if key in run.text:
                    run.text = run.text.replace(key, "" if self.is_empty_value(value) else str(value))
        remaining = {k: v for k, v in placeholders.items() if k in paragraph.text}
        if not remaining:
            return
        text = paragraph.text
        for key, value in remaining.items():
            text = text.replace(key, "" if self.is_empty_value(value) else str(value))
        if paragraph.runs:
            first = paragraph.runs[0]
            paragraph.clear()
            new_run = paragraph.add_run(text)
            new_run.bold = first.bold
            new_run.italic = first.italic
            if first.font.name:
                new_run.font.name = first.font.name
            if first.font.size:
                new_run.font.size = first.font.size

    def _clear_table_cell(self, cell):
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.text = ""
        for child in list(cell._element):
            tag = etree.QName(child).localname if hasattr(child, "tag") else ""
            if tag not in ("p", "tbl", "tcPr"):
                child.getparent().remove(child)

    def remove_partner_blocks(self, doc, partner_prefix):
        for paragraph in list(doc.paragraphs):
            if f"{{{{{partner_prefix}_" in paragraph.text or f"{partner_prefix}_" in paragraph.text:
                paragraph._element.getparent().remove(paragraph._element)
        for table in doc.tables:
            for row in list(table.rows):
                if any(f"{{{{{partner_prefix}_" in cell.text or f"{partner_prefix}_" in cell.text for cell in row.cells):
                    table._element.remove(row._element)

    def clean_empty_partner_sections(self, doc, bid_data):
        shared_fields = []
        signatory_fields = []
        for prefix in ("LEAD", "FIRST", "SECOND"):
            shared_fields += [f"{prefix}_PARTNER_NAME", f"{prefix}_PARTNER_SHORT", f"{prefix}_ADDRESS"]
            signatory_fields += [f"{prefix}_PARTNER_CEO", f"{prefix}_PARTNER_MD1", f"{prefix}_PARTNER_MD2"]
        empty_shared = [field for field in shared_fields if self.is_empty_value(bid_data.get(field))]
        empty_signatories = {field for field in signatory_fields if self.is_empty_value(bid_data.get(field))}
        targets = [f"{{{{{field}}}}}" for field in empty_shared] + empty_shared
        targets += [f"{{{{{field}}}}}" for field in empty_signatories] + list(empty_signatories)
        for paragraph in list(doc.paragraphs):
            if any(target in paragraph.text for target in targets):
                paragraph._element.getparent().remove(paragraph._element)
        for table in doc.tables:
            column_field = {}
            for row in table.rows:
                for index, cell in enumerate(row.cells):
                    for field in signatory_fields:
                        if f"{{{{{field}}}}}" in cell.text:
                            column_field[index] = field
            clear_columns = {index for index, field in column_field.items() if field in empty_signatories}
            for row in list(table.rows):
                if len(row.cells) == 1 and any(target in row.cells[0].text for target in targets):
                    table._element.remove(row._element)
                    continue
                for index, cell in enumerate(row.cells):
                    if index in clear_columns or any(target in cell.text for target in targets):
                        self._clear_table_cell(cell)

    def replace_in_document(self, doc, placeholders):
        paragraphs = list(doc.paragraphs)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    paragraphs.extend(cell.paragraphs)
        for section in doc.sections:
            paragraphs.extend(section.header.paragraphs)
            paragraphs.extend(section.footer.paragraphs)
        for paragraph in paragraphs:
            self.replace_all_in_paragraph(paragraph, placeholders)

    def _all_paragraphs(self, doc):
        yield from doc.paragraphs
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    yield from cell.paragraphs
        for section in doc.sections:
            yield from section.header.paragraphs
            yield from section.footer.paragraphs

    def unresolved_placeholders(self, doc):
        found = set()
        for paragraph in self._all_paragraphs(doc):
            found.update(re.findall(r"\{\{[^{}]+\}\}", paragraph.text))
        return sorted(found)

    def determine_partner_count(self, data):
        if data.get("BID_TYPE") == "Single Bidder":
            if self.is_empty_value(data.get("LEAD_PARTNER_NAME")):
                raise ValueError("Lead partner name is required for single bidder.")
            return 1
        lead = not self.is_empty_value(data.get("LEAD_PARTNER_NAME"))
        first = not self.is_empty_value(data.get("FIRST_PARTNER_NAME"))
        second = not self.is_empty_value(data.get("SECOND_PARTNER_NAME"))
        if not lead:
            raise ValueError("At least the lead partner must be filled to generate the bid.")
        if second and not first:
            raise ValueError("The first partner must be filled before the second partner.")
        return 3 if second else 2 if first else 1

    def select_template(self, partner_count):
        names = {1: "master_template_1.docx", 2: "master_template_2.docx", 3: "master_template_3.docx"}
        selected = names[partner_count]
        if partner_count == 1 and not (self.base_path / "templates" / selected).exists():
            selected = names[2]
        return selected

    def generate_bytes(self, data):
        partner_count = self.determine_partner_count(data)
        template_path = self.base_path / "templates" / self.select_template(partner_count)
        if not template_path.exists():
            raise FileNotFoundError(f"Template not found: {template_path}")
        doc = Document(template_path)
        if partner_count == 1:
            self.remove_partner_blocks(doc, "FIRST")
            self.remove_partner_blocks(doc, "SECOND")
        elif partner_count == 2:
            self.remove_partner_blocks(doc, "SECOND")
        self.clean_empty_partner_sections(doc, data)
        self.replace_in_document(doc, {f"{{{{{key}}}}}": value for key, value in data.items()})
        unresolved = self.unresolved_placeholders(doc)
        if unresolved:
            raise ValueError("The selected template contains unresolved placeholders: " + ", ".join(unresolved))
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        return output.getvalue()

    # A compatibility alias: callers receive bytes instead of a disk path.
    generate = generate_bytes


def _money(value):
    try:
        return f"{float(value):,.2f}"
    except (TypeError, ValueError):
        return "0.00"


def _set_cell(cell, text, bold=False):
    cell.text = str(text or "")
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_after = Pt(0)
        for run in paragraph.runs:
            run.bold = bold
            run.font.size = Pt(8.5)


def _add_table(doc, headers: Iterable[str], rows: Iterable[Iterable[str]]):
    rows = list(rows)
    table = doc.add_table(rows=1, cols=len(list(headers)))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for cell, header in zip(table.rows[0].cells, headers):
        _set_cell(cell, header, bold=True)
    # Repeat the header row when a table spans pages in Word.
    tr_pr = table.rows[0]._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(etree.QName("http://schemas.openxmlformats.org/wordprocessingml/2006/main", "val"), "true")
    tr_pr.append(tbl_header)
    for values in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, values):
            _set_cell(cell, value)
    return table


def _new_form(title, subtitle=None, company_name="", project_name="", ifb_number=""):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)
    heading = doc.add_heading(title, level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if subtitle:
        p = doc.add_paragraph(subtitle)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    metadata = []
    if company_name: metadata.append(("Bidder", company_name))
    if project_name: metadata.append(("Project / tender", project_name))
    if ifb_number: metadata.append(("IFB / contract no.", ifb_number))
    if metadata:
        header_line = "  |  ".join(f"{label}: {value}" for label, value in metadata)
        meta = doc.add_paragraph(header_line)
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in meta.runs:
            run.font.size = Pt(8.5)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("Generated from saved bidder records — verify against the tender-issued SBD.").font.size = Pt(8)
    return doc


def build_fin2_doc(company_name, rows, nrb_indices, current_index, selected_rows, average):
    doc = _new_form("FORM FIN-2", "Average Annual Construction Turnover", company_name=company_name)
    doc.add_paragraph(f"Applicant: {company_name}")
    doc.add_paragraph("The calculations below use the latest published NRB index as the reference index.")
    year_rows = []
    all_jv_rows = []
    for row in rows:
        index = Decimal(str(nrb_indices.get(row.fiscal_year, 0) or 0))
        factor = current_index / index if index else 0
        amount = Decimal(str(row.turnover_amount or 0))
        jv_total = sum((Decimal(str(entry.attributed_amount or 0)) for entry in row.jv_entries), Decimal("0"))
        total = amount + jv_total
        present = total * factor
        year_rows.append((row.fiscal_year, _money(amount), _money(jv_total), _money(total), _money(index), f"{factor:.4f}", _money(present)))
        for entry in row.jv_entries:
            all_jv_rows.append((entry.jv_name, entry.jv_address, entry.vat_number, _money(entry.attributed_amount), f"{float(entry.share_percentage or 0):.2f}%", _money(float(entry.attributed_amount or 0) * float(entry.share_percentage or 0) / 100)))
    _add_table(doc, ["Fiscal year", "Amount", "From JV", "Total", "NRB index", "Factor", "Present value"], year_rows)
    doc.add_heading("Best 3 of the Most Recent 5 Fiscal Years", level=2)
    summary_rows = [(row[0], row[-1]) for row in selected_rows]
    summary_rows.append(("Average Annual Construction Turnover", _money(average)))
    _add_table(doc, ["Selected year", "Escalated amount (NPR)"], summary_rows)
    doc.add_heading("JV Turnover Breakdown", level=2)
    _add_table(doc, ["JV name", "Address", "VAT no.", "Amount", "Share", "Amount-share from JV"], all_jv_rows or [("No JV entries", "", "", "", "", "")])
    return _doc_bytes(doc)


def _experience_block(doc, entry, description_label, description):
    doc.add_heading(f"{entry.contract_id or 'Experience entry'} — {entry.contract_name}", level=2)
    _add_table(doc, ["Field", "Details"], [
        ("Starting / ending month-year", f"{entry.start_month_year} – {entry.end_month_year}"),
        ("Award / completion date", f"{entry.award_date} / {entry.completion_date}"),
        ("Role of bidder", entry.role),
        ("Final Payment Amount (Without VAT) (NRS)", _money(entry.total_contract_amount)),
        ("JV/subcontractor participation", f"{float(entry.participation_percentage or 0):.2f}% / {_money(entry.participation_amount)}"),
        ("Employer", f"{entry.employer_name}; {entry.employer_address}"),
        ("Brief description of works", entry.work_description),
        (description_label, description),
    ])
    items = getattr(entry, "item_quantities", None) or []
    if items:
        doc.add_heading("Key activities and quantities", level=3)
        _add_table(doc, ["Item / activity", "Unit", "Quantity", "From (BS)", "Till (BS)"], [(item.get("item", ""), item.get("unit", ""), item.get("quantity", ""), item.get("from_bs", ""), item.get("till_bs", "")) for item in items])


def _month_year_from_bs(value):
    parts = str(value or "").split("-")
    if len(parts) != 3 or not all(part.isdigit() for part in parts):
        return ""
    year, month = int(parts[0]), int(parts[1])
    names = ("Baisakh", "Jestha", "Ashadh", "Shrawan", "Bhadra", "Ashwin", "Kartik", "Mangsir", "Poush", "Magh", "Falgun", "Chaitra")
    return f"{names[month - 1]} {year}" if 1 <= month <= 12 else ""


def _experience_year(entry):
    """Return the completion year for EXP-1's distinct Year column.

    The reference prompt leaves the meaning of Year ambiguous. Completion
    year is the least surprising interpretation for a completed-project form.
    """
    source = str(entry.completion_date or entry.end_month_year or "")
    matches = re.findall(r"(?<!\d)\d{4}(?!\d)", source)
    return matches[-1] if matches else ""


def _form_heading(doc, title, subtitle):
    heading = doc.add_heading(title, level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph = doc.add_paragraph(subtitle)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _write_exp1(doc, entries):
    rows = []
    for entry in entries:
        start_month = _month_year_from_bs(getattr(entry, "award_date", "")) or getattr(entry, "start_month_year", "")
        end_month = _month_year_from_bs(getattr(entry, "completion_date", "")) or getattr(entry, "end_month_year", "")
        rows.append((start_month, end_month, _experience_year(entry), f"{entry.contract_id} / {entry.contract_name}", f"{entry.employer_name}\\n{entry.employer_address}", entry.work_description, entry.role))
    _add_table(doc, ["Starting month/year", "Ending month/year", "Year", "Contract ID / name", "Employer address", "Brief description", "Role"], rows or [("", "", "", "No entries selected", "", "", "")])
    item_rows = []
    for entry in entries:
        for item in (getattr(entry, "item_quantities", None) or []):
            item_rows.append((item.get("item", ""), item.get("unit", ""), item.get("quantity", ""), item.get("from_bs", ""), item.get("till_bs", ""), f"{entry.contract_id} / {entry.contract_name}"))
    if item_rows:
        doc.add_heading("Supplementary key activity quantities", level=2)
        _add_table(doc, ["Item / activity", "Unit", "Quantity", "From (BS)", "Till (BS)", "Project"], item_rows)


def _write_rolling_summary(doc, item_summary):
    doc.add_heading("Rolling 12-Month Quantity Summary", level=2)
    rows = [(row["item"], row.get("unit") or "", f"{row.get('from_bs', '')} to {row.get('till_bs', '')}", str(row.get("quantity", "")), str(row.get("projects", ""))) for row in item_summary]
    if not rows:
        doc.add_paragraph("No dated item rows are available for a rolling 12-month total.")
        return
    _add_table(doc, ["Item / activity", "Unit", "12-month window (BS)", "Total quantity", "Projects"], rows)


def build_exp1_doc(company_name, entries):
    doc = _new_form("FORM EXP-1", "General Construction Experience", company_name=company_name)
    doc.add_paragraph(f"Bidder: {company_name}")
    _write_exp1(doc, entries)
    return _doc_bytes(doc)


def build_experience_doc(company_name, entries, similarity_entry=None, key_activity_entries=(), item_summary=()):
    """Build a single Word document holding the EXP-1, EXP-2(a) and EXP-2(b) forms."""
    doc = _new_form("Experience Forms", "EXP-1, EXP-2(a) and EXP-2(b)", company_name=company_name)
    doc.add_paragraph(f"Bidder: {company_name}")
    _form_heading(doc, "FORM EXP-1", "General Construction Experience")
    _write_exp1(doc, entries)
    doc.add_page_break()
    _form_heading(doc, "FORM EXP-2(a)", "Specific Construction Experience")
    if similarity_entry is None:
        doc.add_paragraph("No qualifying contract selected.")
    else:
        _experience_block(doc, similarity_entry, "Description of similarity", "")
    doc.add_page_break()
    _form_heading(doc, "FORM EXP-2(b)", "Specific Construction Experience in Key Activities")
    key_activity_entries = list(key_activity_entries)
    for entry, description in key_activity_entries:
        _experience_block(doc, entry, "Production rate description", description)
    if not key_activity_entries:
        doc.add_paragraph("No entries selected.")
    _write_rolling_summary(doc, list(item_summary))
    return _doc_bytes(doc)


def _doc_bytes(doc):
    stream = io.BytesIO()
    doc.save(stream)
    return stream.getvalue()
