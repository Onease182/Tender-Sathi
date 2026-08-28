from decimal import Decimal
from io import BytesIO
from types import SimpleNamespace

from docx import Document
from fastapi.testclient import TestClient

from app import app, bid_data, financial_calculation
from doc_generator import build_exp1_doc, build_experience_doc, build_fin2_doc


def test_landing_page_loads_without_database_configuration(monkeypatch):
    response = TestClient(app).get("/")
    assert response.status_code == 200
    assert "Build bid documents with confidence" in response.text


def test_financial_calculation_selects_best_three_of_latest_five():
    rows = [
        SimpleNamespace(fiscal_year=year, turnover_amount=Decimal(amount), jv_entries=[])
        for year, amount in (("2077/078", "100"), ("2078/079", "500"), ("2079/080", "200"), ("2080/081", "400"), ("2081/082", "300"), ("2082/083", "900"))
    ]
    indices = [SimpleNamespace(fiscal_year=year, index_value=Decimal("100")) for year in ("2077/078", "2078/079", "2079/080", "2080/081", "2081/082", "2082/083")]
    result = financial_calculation(rows, indices)
    assert [item["year"] for item in result["selected"]] == ["2082/083", "2078/079", "2080/081"]
    assert result["average"] == Decimal("600")


def test_single_bidder_normalizes_jv_and_share():
    data = bid_data({
        "BID_TYPE": "Single Bidder",
        "LEAD_PARTNER_NAME": "Lead Builders",
        "LEAD_ADDRESS": "Kathmandu",
        "L_PER": "100.00",
    })
    assert data["JV_NAME"] == "Lead Builders"
    assert data["JV_ADDRESS"] == "Kathmandu"
    assert data["L_PER"] == "100%"
    assert data["AND_CONNECTOR"] == ""


def test_fin2_total_includes_jv_and_present_value_uses_corrected_total():
    entry = SimpleNamespace(fiscal_year="2080/081", turnover_amount=Decimal("41250006.00"), jv_entries=[SimpleNamespace(attributed_amount=Decimal("10389925.20"), share_percentage=Decimal("25"), jv_name="JV Co", jv_address="Kathmandu", vat_number="VAT-1")])
    content = build_fin2_doc("Example Builders", [entry], {"2080/081": Decimal("100")}, Decimal("100"), [("2080/081", "51,639,931.20")], Decimal("51639931.20"))
    table = Document(BytesIO(content)).tables[0]
    assert table.rows[1].cells[3].text == "51,639,931.20"
    assert table.rows[1].cells[6].text == "51,639,931.20"


def test_exp1_year_column_uses_completion_year():
    entry = SimpleNamespace(start_month_year="Jan 2078", end_month_year="Dec 2080", completion_date="2080-12-30", contract_id="C-1", contract_name="Bridge", employer_name="Road Office", employer_address="Pokhara", work_description="Bridge construction", role="Contractor")
    table = Document(BytesIO(build_exp1_doc("Example Builders", [entry]))).tables[0]
    assert [cell.text for cell in table.rows[0].cells[:3]] == ["Starting month/year", "Ending month/year", "Year"]
    assert table.rows[1].cells[2].text == "2080"


def test_experience_document_holds_all_three_forms_and_rolling_summary():
    entry = SimpleNamespace(start_month_year="Jan 2078", end_month_year="Dec 2080", award_date="2078-01-05", completion_date="2080-12-30", contract_id="C-1", contract_name="Bridge", employer_name="Road Office", employer_address="Pokhara", work_description="Bridge construction", role="Contractor", total_contract_amount=Decimal("100"), participation_percentage=Decimal("100"), participation_amount=Decimal("100"), item_quantities=[])
    summary = [{"item": "M20", "unit": "m3", "from_bs": "2080-01-01", "till_bs": "2080-12-30", "quantity": Decimal("17"), "projects": 2}]
    document = Document(BytesIO(build_experience_doc("Example Builders", [entry], entry, [(entry, "")], summary)))
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert "FORM EXP-1" in text and "FORM EXP-2(a)" in text and "FORM EXP-2(b)" in text
    assert "Rolling 12-Month Quantity Summary" in text
    summary_table = document.tables[-1]
    assert [cell.text for cell in summary_table.rows[0].cells] == ["Item / activity", "Unit", "12-month window (BS)", "Total quantity", "Projects"]
    assert [cell.text for cell in summary_table.rows[1].cells] == ["M20", "m3", "2080-01-01 to 2080-12-30", "17", "2"]


def test_experience_document_uses_times_new_roman_and_omits_items_from_exp2a():
    entry = SimpleNamespace(start_month_year="Jan 2078", end_month_year="Dec 2080", award_date="2078-01-05", completion_date="2080-12-30", contract_id="C-1", contract_name="Bridge", employer_name="Road Office", employer_address="Pokhara", work_description="Bridge construction", role="Contractor", total_contract_amount=Decimal("100"), participation_percentage=Decimal("100"), participation_amount=Decimal("100"), item_quantities=[{"item": "M20", "unit": "m3", "quantity": "10", "from_bs": "2080-01-01", "till_bs": "2080-06-01"}])
    document = Document(BytesIO(build_experience_doc("Example Builders", [entry], entry, [(entry, "")], [])))
    # EXP-2(b) keeps its key activity table; EXP-2(a) does not repeat it.
    assert [paragraph.text for paragraph in document.paragraphs].count("Key activities and quantities") == 1
    assert document.styles["Normal"].font.name == "Times New Roman"
    assert min(style.font.size.pt for style in (document.styles[name] for name in ("Normal", "Heading 1", "Heading 2", "Heading 3"))) == 12
    cell_sizes = {run.font.size.pt for table in document.tables for row in table.rows for cell in row.cells for paragraph in cell.paragraphs for run in paragraph.runs if run.font.size}
    assert cell_sizes == {12}


def test_fiscal_year_validation_rejects_typos():
    from app import fy_sort
    assert fy_sort("2080/081") == 2080
    import pytest
    with pytest.raises(ValueError):
        fy_sort("2080-081")


def test_safe_next_rejects_external_redirects():
    from app import safe_next
    assert safe_next("/dashboard?section=generate") == "/dashboard?section=generate"
    assert safe_next("https://example.com/account") == "/dashboard"
    assert safe_next("//example.com/account") == "/dashboard"
